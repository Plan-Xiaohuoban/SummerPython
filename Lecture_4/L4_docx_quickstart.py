from docx import Document

document = Document()

paragraph = document.add_paragraph("Lorem ipsum dolor sit amet.")


prior_paragraph = paragraph.insert_paragraph_before("Lorem ipsum")


document.add_heading("The REAL meaning of the universe")

document.add_heading("The role of dolphins", level=2)


document.add_page_break()


table = document.add_table(rows=2, cols=2)


cell = table.cell(0, 1)


cell.text = "parrot, possibly dead"

row = table.rows[1]
row.cells[0].text = "Foo bar to you."
row.cells[1].text = "And a hearty foo bar to you too sir!"

for row in table.rows:
    for cell in row.cells:
        print(cell.text)
row = table.add_row()

# get table data -------------
items = (
    (7, "1024", "Plush kittens"),
    (3, "2042", "Furbees"),
    (1, "1288", "French Poodle Collars, Deluxe"),
)

# add table ------------------
table = document.add_table(1, 3)

# populate header row --------
heading_cells = table.rows[0].cells
heading_cells[0].text = "Qty"
heading_cells[1].text = "SKU"
heading_cells[2].text = "Description"


row_count = len(table.rows)
col_count = len(table.columns)

document.save("test.docx")
