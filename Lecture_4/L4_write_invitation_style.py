import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH  # 设置对齐方式
from docx.oxml.ns import qn  # 设置字体
from docx.shared import Cm, Inches, Pt, RGBColor  # 用来设置字体大小、缩进; 设置字体的颜色

from L4_info import *

# 创建文档对象
doc = docx.Document()

# 空行
for i in range(5):
    doc.add_paragraph("")

heading = doc.add_paragraph("")  # 创建段落对象，当作标题
heading.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 设置为居中对齐
heading.paragraph_format.space_before = Pt(0)  # 设置段前间距 0 磅
heading.paragraph_format.space_after = Pt(0)  # 设置段后 0 磅
heading.paragraph_format.line_spacing = 1  # 设置行间距为 1 倍
heading.paragraph_format.left_indent = Inches(0)  # 设置左缩进 0 英寸
heading.paragraph_format.right_indent = Inches(0)  # 设置右缩进 0 英寸


run = heading.add_run(title)
run.font.name = u"黑体"  # 设置字体
run._element.rPr.rFonts.set(qn("w:eastAsia"), u"黑体")  # 设置为字体，和上边的保持一致
run.font.size = Pt(20)  # 设置文字的大小为20磅
run.font.color.rgb = RGBColor(0, 0, 0)  # 设置颜色为黑色


# 设置正文格式
doc.styles["Normal"].font.size = Pt(14)
doc.styles["Normal"].font.name = u"仿宋"  # 设置当文字是西文时的字体
doc.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), u"仿宋")  # 设置当文字是中文时的字体


# 添加内容
for i in range(2):
    doc.add_paragraph("")

doc.add_paragraph(receiver)

paragraph = doc.add_paragraph(text)
paragraph.paragraph_format.first_line_indent = Cm(1)  # 设置左缩进 0 英寸


paragraph = doc.add_paragraph(sender)
paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

paragraph = doc.add_paragraph(date)
paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

doc.save("Output/社会实践邀请函_仿制_含格式.docx")
