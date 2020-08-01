import docx
from L4_info import *  # 从L4_info文件导入写好的文本

from docx.oxml.ns import qn  # 设置字体
from docx.shared import Cm, Inches, Pt, RGBColor  # 用来设置字体大小、缩进; 设置字体的颜色


# 创建文档对象
doc = docx.Document()

# 添加标题
doc.add_heading(title, level=1)


# 设置正文格式
doc.styles["Normal"].font.size = Pt(14)
doc.styles["Normal"].font.name = u"仿宋"  # 设置当文字是西文时的字体
doc.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), u"仿宋")  # 设置当文字是中文时的字体



# 添加段落
doc.add_paragraph(receiver)
doc.add_paragraph(text)
doc.add_paragraph(sender)
doc.add_paragraph(date)

# 保存文档（保存时请确认该文档是关闭的，否则会出权限错误）
doc.save("Output/社会实践邀请函_仿制.docx")

