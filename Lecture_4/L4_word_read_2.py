# 若无法导入，则pip install python-docx
import docx

file = docx.Document("Data_samples/社会实践邀请函范例.docx")  # 创建文档对象
print(file.paragraphs)
print("段落数:" + str(len(file.paragraphs)))  # 读取文档的段落数量

# 输出段落编号及段落内容的前三十个字
for i, paragraph in enumerate(file.paragraphs):
    print("第" + str(i) + "段：" + paragraph.text[:30])


# 输出每一段的内容
# for para in file.paragraphs:
#     print(para.text)
